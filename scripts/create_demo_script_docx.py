"""Generate the Fincast Demo Video Script as a professional Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    if level == 1:
        h.font.size = Pt(26)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(16)
        h.font.bold = True
    elif level == 3:
        h.font.size = Pt(13)
        h.font.bold = True

# Quote style
quote_style = doc.styles.add_style("ScriptQuote", 1)  # paragraph
quote_style.font.name = "Georgia"
quote_style.font.size = Pt(11.5)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
quote_style.paragraph_format.left_indent = Cm(1.2)
quote_style.paragraph_format.space_after = Pt(8)
quote_style.paragraph_format.line_spacing = 1.3

# Stage direction style
stage_style = doc.styles.add_style("StageDirection", 1)
stage_style.font.name = "Calibri"
stage_style.font.size = Pt(10)
stage_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
stage_style.paragraph_format.left_indent = Cm(1.2)
stage_style.paragraph_format.space_after = Pt(4)

# ── Helpers ─────────────────────────────────────────────────────────────────

ORANGE = RGBColor(0xFF, 0x6B, 0x35)
CHARCOAL = RGBColor(0x0A, 0x0A, 0x0A)
GRAY = RGBColor(0x66, 0x66, 0x66)


def add_narration(text, bold_phrases=None):
    """Add a narration paragraph in Georgia italic with optional bold phrases."""
    p = doc.add_paragraph(style="ScriptQuote")
    if bold_phrases is None:
        bold_phrases = []

    remaining = text
    for phrase in bold_phrases:
        if phrase in remaining:
            before, remaining = remaining.split(phrase, 1)
            if before:
                p.add_run(before)
            run = p.add_run(phrase)
            run.bold = True
    if remaining:
        p.add_run(remaining)
    return p


def add_stage(text):
    """Add a stage direction line."""
    p = doc.add_paragraph(style="StageDirection")
    run = p.add_run(f"[{text}]")
    return p


def add_divider():
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_timing_badge(section_num, title, start, end, duration):
    """Add a section header with timing info."""
    doc.add_heading(f"Section {section_num}: {title}", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"TIME: {start} - {end}  |  DURATION: {duration}")
    run.font.size = Pt(9)
    run.font.color.rgb = ORANGE
    run.font.bold = True
    run.font.name = "Calibri"


def set_table_style(table):
    """Style a table with orange header."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"
            if row_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "FF6B35")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True


# ── Page Setup ──────────────────────────────────────────────────────────────

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Title Page ──────────────────────────────────────────────────────────────

# Spacer
for _ in range(6):
    doc.add_paragraph()

title = doc.add_heading("FINCAST", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(48)
    run.font.color.rgb = ORANGE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("5-Minute Demo Video Script")
run.font.size = Pt(20)
run.font.color.rgb = CHARCOAL
run.font.name = "Calibri"

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("AI-Powered Podcast Platform for Finance Professionals")
run.font.size = Pt(13)
run.font.color.rgb = GRAY
run.font.italic = True
run.font.name = "Georgia"

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Total Duration: 5:00  |  ~893 Words  |  7 Sections")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_page_break()

# ── Structure Overview ──────────────────────────────────────────────────────

doc.add_heading("Script Structure", level=1)

table = doc.add_table(rows=8, cols=4)
headers = ["Section", "Time", "Duration", "Purpose"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ("1. Hook & Problem", "0:00-0:40", "40s", "Why this matters"),
    ("2. Solution Overview", "0:40-1:20", "40s", "What Fincast is"),
    ("3. Home & Discovery", "1:20-2:00", "40s", "Browse, play, star, search"),
    ("4. AI News Podcast", "2:00-3:20", "80s", 'The "wow" feature (live)'),
    ("5. Macro Tracker & Chat", "3:20-4:10", "50s", "Intelligence layer"),
    ("6. Architecture", "4:10-4:40", "30s", "How it works"),
    ("7. Closing", "4:40-5:00", "20s", "Impact & future"),
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

set_table_style(table)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Narrative technique: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Opens and closes with "500 hours" (bookend structure). Section 4 is paced as a live magic trick. Architecture closes with "Zero mocks." Final line: "This is the future of market intelligence."')
run.font.size = Pt(10)
run.font.color.rgb = GRAY

add_divider()

# ── Section 1 ───────────────────────────────────────────────────────────────

add_timing_badge(1, "Hook & Problem", "0:00", "0:40", "40 seconds")

add_stage("SCREEN: Dark/black. Fincast logo fades in at 0:03. Optional news montage 0:10-0:35. Cut to app at 0:35.")

add_narration(
    '"Five hundred hours a year. That\'s how long the average fund manager spends just reading the news. Not analyzing. Not trading. Reading.\n\n'
    "Hundreds of articles, reports, data points \u2014 every single morning. And most of it? Noise.\n\n"
    'What if you could skip the noise and get the signal \u2014 as a podcast you listen to on your commute? That\'s Fincast."',
    bold_phrases=["Five hundred hours a year.", "Noise.", "Fincast."],
)

add_divider()

# ── Section 2 ───────────────────────────────────────────────────────────────

add_timing_badge(2, "Solution Overview", "0:40", "1:20", "40 seconds")

add_stage("SCREEN: Fincast home page, logged in, three-column layout. HOVER ONLY \u2014 no clicks. Move cursor slowly across podcast grid, then point to right sidebar Trending Topics.")

add_narration(
    '"Fincast is an AI-powered podcast platform built for finance professionals. Three capabilities:\n\n'
    "First \u2014 it takes breaking financial news and turns it into broadcast-quality audio briefings. Sixty seconds, start to finish. AI does the research, writes the script, generates the voice.\n\n"
    "Second \u2014 it tracks macro-economic themes using Google Trends momentum data. You see what\u2019s heating up before the market catches on.\n\n"
    "Third \u2014 an AI assistant that understands macro context and gives you actionable answers, not summaries.\n\n"
    'Let me show you."',
    bold_phrases=["Three capabilities:", "First", "Second", "Third"],
)

add_divider()

# ── Section 3 ───────────────────────────────────────────────────────────────

add_timing_badge(3, "Home & Discovery", "1:20", "2:00", "40 seconds")

add_stage("ACTIONS: Click podcast card \u2192 detail page. Click Play \u2192 sticky player appears. Click Star \u2192 yellow fill. Navigate to Favorites. Navigate to Discover \u2192 type search query. Pause player before Section 4.")

add_narration(
    '"Home page. Podcasts ranked by our heat scoring algorithm \u2014 I\'ll break that down shortly.\n\n'
    "Each podcast has a full detail page: transcript, AI-generated cover art, linked macro themes. Hit play \u2014 the sticky player follows you everywhere in the app.\n\n"
    "Star your favorites. They\u2019re saved here for quick access.\n\n"
    'And discover \u2014 full-text search across every title, description, and transcript on the platform. Find anything instantly."',
    bold_phrases=["heat scoring algorithm"],
)

add_divider()

# ── Section 4 ───────────────────────────────────────────────────────────────

add_timing_badge(4, "AI News Podcast Creation", "2:00", "3:20", "80 seconds")

p = doc.add_paragraph()
run = p.add_run("THIS IS THE SHOWSTOPPER. Pace it like a live magic trick.")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = ORANGE

add_stage("Click 'News Podcast' in sidebar \u2192 5-step wizard loads. Follow steps below.")

add_narration(
    '"Alright. This is the core of Fincast. I\'m going to create a professional news podcast from scratch \u2014 live, right now."',
    bold_phrases=[],
)

# Step 1
doc.add_heading("Step 1: Pick a Topic", level=3)
add_stage("Click 'US Inflation' topic card (look for red 'Hot' badge). Wait 5-10s for articles to load.")

add_narration(
    '"Step one: pick a topic. I\'ll go with US Inflation \u2014 notice it\'s flagged \'Hot,\' heat score 2.1. The moment I select it, GPT-4.1-mini fires a live web search and pulls the top trending articles. These aren\'t cached. These are real articles, fetched seconds ago."',
    bold_phrases=["Hot"],
)

# Step 2
doc.add_heading("Step 2: Review Articles", level=3)
add_stage("Scroll through articles. Optionally toggle one off/on. Click 'Generate Script'.")

add_narration(
    '"Step two: I pick the stories I want covered. Three articles selected. Next."',
)

# Step 3
doc.add_heading("Step 3: Generate Script", level=3)
add_stage("Change tone to 'Professional'. Scroll through generated script. Point to word count / duration stats. Click 'Continue to Audio'.")

add_narration(
    '"Step three: generate the script. Professional tone, medium length. Watch this.\n\n'
    'A complete podcast script \u2014 sourced, structured, ready to record. I can edit anything. Notice the live word count and estimated duration."',
    bold_phrases=["Watch this."],
)

# Step 4
doc.add_heading("Step 4: Audio & Cover Art", level=3)
add_stage("Select voice 'Nova'. Click 'Generate Audio' (wait 10-30s). While waiting, scroll to cover art and click 'Generate Thumbnail' (wait 10-20s). Narrate during both waits.")

add_narration(
    '"Step four: this is where it gets interesting. I pick a voice \u2014 Nova \u2014 and hit generate. OpenAI\'s TTS engine converts the entire script to speech. Our backend splits at sentence boundaries to handle any script length, then stitches the audio seamlessly.\n\n'
    "While that renders, I\u2019ll generate the cover art. GPT reads the script, writes a DALL-E 3 prompt based on the actual story \u2014 not generic podcast art.\n\n"
    'Audio ready. Art ready."',
    bold_phrases=["this is where it gets interesting."],
)

# Step 5
doc.add_heading("Step 5: Publish", level=3)
add_stage("Click 'Review & Publish' \u2192 review screen. Click 'Publish News Podcast'. Redirects to home page with new podcast in grid.")

add_narration(
    '"Step five: publish.\n\n'
    'Done. Topic to published podcast. The user touched five buttons."',
    bold_phrases=["The user touched five buttons."],
)

# Fallbacks
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Fallback plans:")
run.bold = True
run.font.size = Pt(10)

fallbacks = [
    ("Article fetch >10s:", '"The system is querying live news sources in real-time \u2014 no cached data."'),
    ("Audio >45s:", '"For longer scripts, this can take up to a minute." Navigate to pre-created backup.'),
    ("DALL-E fails:", 'Click "Upload custom image" tab as alternative.'),
    ("Critical backup:", "Pre-create one podcast with the same topic before recording."),
]
for label, desc in fallbacks:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"\u2022 {label} ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

add_divider()

# ── Section 5 ───────────────────────────────────────────────────────────────

add_timing_badge(5, "Macro Tracker & AI Chat", "3:20", "4:10", "50 seconds")

add_stage("Click theme badge on podcast detail \u2192 topic page. Point to metrics grid. Scroll to Latest Developments. Open chatbot (orange button, bottom-right). Type question. Wait for response.")

add_narration(
    '"Every published podcast gets auto-tagged with macro-economic themes by GPT. These badges link to our Macro Economics Tracker.\n\n'
    "This is where Fincast becomes an intelligence tool. Each theme has a heat score computed from Google Trends \u2014 but not raw search volume. We measure momentum: the rate of change. Above 1.15 means it\u2019s accelerating. Below 0.88, it\u2019s cooling. This tells you what\u2019s moving before consensus forms.\n\n"
    "You see the score, mention count, daily change, and a 7-day sparkline. Below that, auto-generated summaries of latest developments with source links \u2014 GPT with live web search.\n\n"
    'Now the AI assistant. It\'s context-aware \u2014 it already knows I\'m looking at US Inflation."',
    bold_phrases=["momentum: the rate of change", "before consensus forms"],
)

add_stage('TYPE in chatbot: "What are the risk implications for bond portfolios?" \u2192 Send. Wait 3-8s.')

add_narration(
    '"I\'ll ask: \'What are the risk implications for bond portfolios?\'\n\n'
    'Watch the response. It draws from theme data, recent developments, and financial domain knowledge. This isn\'t a generic chatbot. It gives you actionable analysis."',
    bold_phrases=["actionable analysis"],
)

add_divider()

# ── Section 6 ───────────────────────────────────────────────────────────────

add_timing_badge(6, "Technical Architecture", "4:10", "4:40", "30 seconds")

add_stage("OPTION A: Switch to pre-made architecture diagram. OPTION B: Narrate over the app (no clicks).")

add_narration(
    '"Under the hood: Next.js 16 with React 19 on the frontend. Convex as the serverless backend \u2014 real-time subscriptions out of the box. Data changes, UI updates instantly. No polling.\n\n'
    "OpenAI powers four AI pipelines: GPT-4.1-mini for news fetching, script generation, theme analysis, and the chatbot. TTS-1 for voice synthesis with sentence-aware chunking. DALL-E 3 for cover art.\n\n"
    "Google Trends feeds a cron job every three hours, computing momentum-based heat scores across all tracked themes. Clerk handles auth. Resend handles email.\n\n"
    'Everything you saw is real data, real APIs, real-time. Zero mocks."',
    bold_phrases=["Zero mocks."],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run("SHORT VERSION (if running long): ")
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = ORANGE
run = p.add_run('"Next.js 16, Convex for real-time serverless, OpenAI for all AI pipelines, Google Trends for momentum scoring. Everything is real data, real APIs, real-time."')
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.italic = True

add_divider()

# ── Section 7 ───────────────────────────────────────────────────────────────

add_timing_badge(7, "Closing", "4:40", "5:00", "20 seconds")

add_stage("Navigate to home page (if not already). Hold on full three-column layout. Fade to closing slide with logo + team info.")

add_narration(
    '"Fund managers spend five hundred hours a year reading news. Fincast gives them that time back \u2014 with AI audio briefings, real-time trend intelligence, and a financial AI assistant, all in one platform.\n\n'
    "We built this in seven sessions. The codebase is open source.\n\n"
    'This is the future of market intelligence. Thank you."',
    bold_phrases=["five hundred hours", "This is the future of market intelligence."],
)

doc.add_page_break()

# ── Word Count Summary ──────────────────────────────────────────────────────

doc.add_heading("Word Count Summary", level=1)

table = doc.add_table(rows=9, cols=3)
headers = ["Section", "Words", "Duration"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

counts = [
    ("1. Hook & Problem", "~82", "40s"),
    ("2. Solution Overview", "~100", "40s"),
    ("3. Home & Discovery", "~88", "40s"),
    ("4. AI News Podcast Creation", "~248", "80s"),
    ("5. Macro Tracker & Chat", "~195", "50s"),
    ("6. Architecture", "~118", "30s"),
    ("7. Closing", "~62", "20s"),
    ("TOTAL", "~893", "5:00"),
]
for row_idx, row_data in enumerate(counts):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        if row_idx == len(counts) - 1:
            for run in cell.paragraphs[0].runs:
                run.bold = True

set_table_style(table)

add_divider()

# ── Architecture Diagram Reference ──────────────────────────────────────────

doc.add_heading("Architecture Diagram Reference", level=1)

p = doc.add_paragraph()
run = p.add_run("Use this to create a slide for Section 6 (Excalidraw, Figma, or PowerPoint):")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_paragraph()

# Architecture as a clean text block
arch_lines = [
    "FRONTEND     Next.js 16 + React 19 + Tailwind CSS + shadcn/ui",
    "             Pages: Home, Discover, Create, Detail, Topics, Favorites, Profile",
    "             Components: Sticky Player, AI Chatbot, ThemeBadge, Sparklines",
    "",
    "AUTH         Clerk (OAuth/email) \u2192 Svix webhooks \u2192 Convex user sync",
    "",
    "BACKEND      Convex (serverless, real-time subscriptions)",
    "             Queries (cached, auto-subscribe) | Mutations (auth-gated) | Actions (AI/external)",
    "",
    "AI LAYER     GPT-4.1-mini: 6 pipelines (news fetch, script gen, theme tag, chatbot, image prompt, summary)",
    "             TTS-1: voice synthesis with sentence-aware 4096-char chunking",
    "             DALL-E 3: contextual cover art (1024x1024)",
    "",
    "DATA         Convex DB: 5 tables (podcasts, users, macroThemes, themeMentions, favorites)",
    "             13 indexes + 3 full-text search indexes | File storage (MP3/PNG)",
    "",
    "EXTERNAL     Google Trends API (3hr cron \u2192 momentum scoring) | Resend (email delivery)",
]

for line in arch_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    if line and not line.startswith(" "):
        run.font.bold = True
        run.font.color.rgb = ORANGE

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Key stats for judges: ")
run.bold = True
run.font.size = Pt(10)

stats = [
    "8 AI calls per podcast (6 GPT pipelines + TTS + DALL-E)",
    "5 database tables, 13 indexes, 3 full-text search indexes",
    "Real-time UI via Convex subscriptions (zero polling)",
    "Momentum-based scoring (rate of change, not absolute volume)",
    "Zero backend infrastructure to manage",
]
for stat in stats:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"\u2022 {stat}")
    run.font.size = Pt(10)

# ── Save ────────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs",
    "Fincast_Demo_Script.docx",
)
doc.save(output_path)
print(f"Saved to: {output_path}")
